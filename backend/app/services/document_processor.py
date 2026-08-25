import io
import re
import unicodedata
from typing import Tuple, Optional
from pypdf import PdfReader
import docx


class DocumentProcessingError(Exception):
    pass


class DocumentProcessor:
    """
    Handles robust extraction, cleaning, and preprocessing of multiple file formats (.txt, .pdf, .docx, .md).
    """

    SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".doc", ".md", ".markdown"}
    MAX_FILE_SIZE_BYTES = 25 * 1024 * 1024  # 25 MB

    @staticmethod
    def clean_and_normalize_text(text: str) -> str:
        """
        Cleans and normalizes extracted text:
        - Normalizes Unicode characters (quotes, dashes, spaces)
        - Normalizes whitespace while preserving paragraphs
        - Strips control characters
        """
        if not text:
            return ""

        # Normalize unicode NFKC
        text = unicodedata.normalize("NFKC", text)

        # Replace non-standard whitespace characters with normal spaces
        text = re.sub(r"[\r\f\v]", "\n", text)
        
        # Replace non-breaking spaces and exotic spaces with standard space
        text = text.replace("\u00a0", " ").replace("\u200b", "")

        # Normalize quotes and dashes
        text = text.replace("“", '"').replace("”", '"').replace("’", "'").replace("‘", "'")
        text = text.replace("–", "-").replace("—", "-")

        # Normalize lines: remove trailing spaces from each line
        lines = [line.strip() for line in text.split("\n")]

        # Group into paragraphs (coalesce multiple blank lines into max 2 newlines)
        cleaned_paragraphs = []
        current_para = []

        for line in lines:
            if line:
                current_para.append(line)
            else:
                if current_para:
                    # Join lines within the same paragraph with single space
                    para_text = " ".join(current_para)
                    # Clean multiple consecutive internal spaces
                    para_text = re.sub(r"[ \t]+", " ", para_text)
                    cleaned_paragraphs.append(para_text)
                    current_para = []

        if current_para:
            para_text = " ".join(current_para)
            para_text = re.sub(r"[ \t]+", " ", para_text)
            cleaned_paragraphs.append(para_text)

        cleaned_text = "\n\n".join(cleaned_paragraphs).strip()
        return cleaned_text

    @classmethod
    def extract_from_bytes(cls, file_bytes: bytes, filename: str) -> Tuple[str, dict]:
        """
        Extracts clean text and metadata from raw bytes based on filename extension.
        Returns (extracted_text, metadata_dict).
        """
        if len(file_bytes) > cls.MAX_FILE_SIZE_BYTES:
            raise DocumentProcessingError(
                f"File size ({len(file_bytes) / (1024*1024):.2f} MB) exceeds maximum allowed size (25 MB)."
            )

        lower_name = filename.lower()
        extension = ""
        for ext in cls.SUPPORTED_EXTENSIONS:
            if lower_name.endswith(ext):
                extension = ext
                break

        if not extension:
            # Check default plain text
            extension = ".txt"

        metadata = {
            "filename": filename,
            "extension": extension,
            "size_bytes": len(file_bytes),
            "pages_or_sections": 1,
        }

        try:
            if extension in {".txt", ".md", ".markdown"}:
                text, meta = cls._extract_plain_text(file_bytes)
            elif extension == ".pdf":
                text, meta = cls._extract_pdf(file_bytes)
            elif extension in {".docx", ".doc"}:
                text, meta = cls._extract_docx(file_bytes)
            else:
                text, meta = cls._extract_plain_text(file_bytes)

            metadata.update(meta)
            normalized = cls.clean_and_normalize_text(text)

            if not normalized or len(normalized.strip()) < 5:
                raise DocumentProcessingError(
                    "No readable text could be extracted from the document. The file may be empty, image-only/scanned, or corrupted."
                )

            return normalized, metadata

        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(f"Failed to extract text from '{filename}': {str(e)}")

    @staticmethod
    def _extract_plain_text(file_bytes: bytes) -> Tuple[str, dict]:
        # Try UTF-8, then UTF-16, latin-1, cp1252
        for enc in ["utf-8", "utf-8-sig", "latin-1", "cp1252", "utf-16"]:
            try:
                decoded = file_bytes.decode(enc)
                return decoded, {"encoding": enc}
            except (UnicodeDecodeError, LookupError):
                continue
        # Fallback with replacement
        return file_bytes.decode("utf-8", errors="replace"), {"encoding": "utf-8-lossy"}

    @staticmethod
    def _extract_pdf(file_bytes: bytes) -> Tuple[str, dict]:
        pdf_stream = io.BytesIO(file_bytes)
        try:
            reader = PdfReader(pdf_stream)
        except Exception as e:
            raise DocumentProcessingError(f"Corrupted or password-protected PDF document: {str(e)}")

        num_pages = len(reader.pages)
        if num_pages == 0:
            raise DocumentProcessingError("The PDF document contains 0 pages.")

        extracted_pages = []
        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    extracted_pages.append(page_text.strip())
            except Exception:
                continue

        full_text = "\n\n".join(extracted_pages)
        return full_text, {"pages_or_sections": num_pages, "extracted_pages": len(extracted_pages)}

    @staticmethod
    def _extract_docx(file_bytes: bytes) -> Tuple[str, dict]:
        docx_stream = io.BytesIO(file_bytes)
        try:
            doc = docx.Document(docx_stream)
        except Exception as e:
            raise DocumentProcessingError(f"Corrupted or invalid DOCX document: {str(e)}")

        paragraphs = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                paragraphs.append(p.text.strip())

        # Also extract table cells if any
        table_rows_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_rows_text.append(row_text)

        all_text = "\n\n".join(paragraphs)
        if table_rows_text:
            all_text += "\n\n" + "\n".join(table_rows_text)

        return all_text, {"paragraphs_extracted": len(paragraphs), "tables_extracted": len(doc.tables)}
